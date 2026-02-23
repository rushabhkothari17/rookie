"""Article document generation utilities: HTML → PDF and HTML → DOCX."""
from __future__ import annotations

import io
import re
from html.parser import HTMLParser
from typing import List, Tuple


# ─── HTML Parser ────────────────────────────────────────────────────────────────

class _Block:
    """Represents a structured block from HTML."""
    def __init__(self, kind: str, text: str = "", level: int = 0, bold: bool = False, italic: bool = False):
        self.kind = kind  # "heading", "paragraph", "list_item", "hr", "spacer"
        self.text = text
        self.level = level
        self.bold = bold
        self.italic = italic


class _HTMLToBlocks(HTMLParser):
    def __init__(self):
        super().__init__()
        self._blocks: List[_Block] = []
        self._current_text = ""
        self._current_kind = "paragraph"
        self._current_level = 0
        self._in_list = False
        self._bold = False
        self._italic = False
        self._skip = False

    def handle_starttag(self, tag: str, attrs):
        tag = tag.lower()
        if tag in ("h1", "h2", "h3", "h4"):
            self._flush()
            level = int(tag[1])
            self._current_kind = "heading"
            self._current_level = level
        elif tag == "p":
            self._flush()
            self._current_kind = "paragraph"
        elif tag in ("ul", "ol"):
            self._flush()
            self._in_list = True
        elif tag == "li":
            self._flush()
            self._current_kind = "list_item"
        elif tag in ("strong", "b"):
            self._bold = True
        elif tag in ("em", "i"):
            self._italic = True
        elif tag == "hr":
            self._flush()
            self._blocks.append(_Block("hr"))
        elif tag == "br":
            self._current_text += "\n"
        elif tag in ("img", "figure"):
            self._skip = True

    def handle_endtag(self, tag: str):
        tag = tag.lower()
        if tag in ("h1", "h2", "h3", "h4", "p", "li"):
            self._flush()
            self._current_kind = "paragraph"
            self._current_level = 0
        elif tag in ("ul", "ol"):
            self._in_list = False
        elif tag in ("strong", "b"):
            self._bold = False
        elif tag in ("em", "i"):
            self._italic = False
        elif tag in ("img", "figure"):
            self._skip = False

    def handle_data(self, data: str):
        if self._skip:
            return
        self._current_text += data

    def _flush(self):
        text = self._current_text.strip()
        if text:
            self._blocks.append(_Block(
                kind=self._current_kind,
                text=text,
                level=self._current_level,
                bold=self._bold,
                italic=self._italic,
            ))
        self._current_text = ""

    def get_blocks(self) -> List[_Block]:
        self._flush()
        return self._blocks


def parse_html(html: str) -> List[_Block]:
    parser = _HTMLToBlocks()
    # Remove script/style tags
    html = re.sub(r"<(script|style)[^>]*>.*?</(script|style)>", "", html, flags=re.DOTALL | re.IGNORECASE)
    parser.feed(html)
    return parser.get_blocks()


# ─── PDF Generation ──────────────────────────────────────────────────────────────

def generate_pdf(title: str, author: str, created_at: str, updated_at: str, html_content: str) -> bytes:
    """Generate a PDF from HTML article content using reportlab."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch, cm
    from reportlab.lib.colors import HexColor, black, grey
    from reportlab.lib.enums import TA_LEFT, TA_CENTER
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
    from reportlab.platypus import ListFlowable, ListItem

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=1.5 * cm, rightMargin=1.5 * cm,
        topMargin=2 * cm, bottomMargin=2 * cm,
    )

    styles = getSampleStyleSheet()
    dark = HexColor("#0f172a")
    muted = HexColor("#64748b")
    accent = HexColor("#dc2626")

    style_meta = ParagraphStyle("meta", parent=styles["Normal"], fontSize=8, textColor=muted, spaceAfter=2)
    style_h1 = ParagraphStyle("H1", parent=styles["Heading1"], fontSize=22, textColor=dark, spaceAfter=6, spaceBefore=4)
    style_h2 = ParagraphStyle("H2", parent=styles["Heading2"], fontSize=16, textColor=dark, spaceAfter=4, spaceBefore=10)
    style_h3 = ParagraphStyle("H3", parent=styles["Heading3"], fontSize=13, textColor=dark, spaceAfter=3, spaceBefore=8)
    style_h4 = ParagraphStyle("H4", parent=styles["Heading4"], fontSize=11, textColor=dark, spaceAfter=2, spaceBefore=6)
    style_body = ParagraphStyle("body", parent=styles["Normal"], fontSize=10, textColor=dark, leading=16, spaceAfter=6)
    style_li = ParagraphStyle("li", parent=styles["Normal"], fontSize=10, textColor=dark, leading=14, leftIndent=16, spaceAfter=3)

    heading_styles = {1: style_h1, 2: style_h2, 3: style_h3, 4: style_h4}

    story = []

    # Header
    story.append(Paragraph(f"<font color='#{accent.hexval()[2:]}'>—</font> {title}", style_h1))
    story.append(Paragraph(f"Created: {created_at[:10] if created_at else '—'} &nbsp;|&nbsp; Last updated: {updated_at[:10] if updated_at else '—'} &nbsp;|&nbsp; Author: {author}", style_meta))
    story.append(HRFlowable(width="100%", thickness=1, color=HexColor("#e2e8f0"), spaceAfter=12, spaceBefore=6))

    blocks = parse_html(html_content)
    for block in blocks:
        if block.kind == "heading":
            style = heading_styles.get(block.level, style_h2)
            story.append(Paragraph(block.text, style))
        elif block.kind == "paragraph":
            story.append(Paragraph(block.text, style_body))
        elif block.kind == "list_item":
            story.append(Paragraph(f"• {block.text}", style_li))
        elif block.kind == "hr":
            story.append(HRFlowable(width="100%", thickness=1, color=HexColor("#e2e8f0"), spaceAfter=8, spaceBefore=8))
        elif block.kind == "spacer":
            story.append(Spacer(1, 6))

    doc.build(story)
    return buf.getvalue()


# ─── DOCX Generation ─────────────────────────────────────────────────────────────

def generate_docx(title: str, author: str, created_at: str, updated_at: str, html_content: str) -> bytes:
    """Generate a DOCX from HTML article content using python-docx."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    document = Document()

    # Page margins
    for section in document.sections:
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)

    # Title
    title_para = document.add_heading(title, 0)
    title_para.runs[0].font.color.rgb = RGBColor(0x0f, 0x17, 0x2a)

    # Metadata
    meta = document.add_paragraph()
    meta.add_run(f"Created: {created_at[:10] if created_at else '—'}  |  "
                 f"Last updated: {updated_at[:10] if updated_at else '—'}  |  Author: {author}")
    meta.runs[0].font.size = Pt(8)
    meta.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8b)

    # Divider (simulate with empty line)
    document.add_paragraph()

    blocks = parse_html(html_content)
    for block in blocks:
        if block.kind == "heading":
            lvl = min(block.level, 4)
            p = document.add_heading(block.text, level=lvl)
        elif block.kind == "paragraph":
            p = document.add_paragraph(block.text)
            p.style = document.styles["Normal"]
        elif block.kind == "list_item":
            p = document.add_paragraph(block.text, style="List Bullet")
        elif block.kind == "hr":
            document.add_paragraph()
        # spacer — skip

    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()
